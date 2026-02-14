#!/usr/bin/env python3

from io import StringIO
import pandas as pd
import sys
import os
import tqdm
import altair as alt
from pathlib import Path
from textwrap import wrap

# Define constants
BUDGET_COLUMN = 'Budget (MEUR)'
from docx import Document
from docx.shared import Pt

# Import the bar chart function from our module
from nlesc_ser2026_plots.bar_charts import create_pie_chart, create_sorted_bar_chart, create_survey_chart, create_table_heatmap, create_yearly_stacked_bar_chart, create_yearly_multi_bar_chart, create_yearly_stacked_bar_line_chart
from nlesc_ser2026_plots.geo_charts import plot_netherlands_with_institutions
from nlesc_ser2026_plots.reference_lists import create_refstrings_list
from nlesc_ser2026_plots.line_charts import create_yearly_multi_line_chart
import argparse

def parse_arguments():
    """Parse command-line arguments."""
    parser = argparse.ArgumentParser(description="Generate charts from CSV files in a specified directory.")
    parser.add_argument("input_dir", type=str, help="Path to the directory containing the CSV files.")
    parser.add_argument("--format", type=str, default="png", choices=["html", "png", "svg", "pdf"], help="Output format for the charts.")
    parser.add_argument("--output_dir", type=str, default=".", help="Directory to save the generated charts.")
    return parser.parse_args()


def list_references(
        output_file: Path, 
        df: pd.DataFrame, 
        offset: int = 1    
    ) -> int:
    """
    List references from a DataFrame to a DOCX file.
    Args:
        output_file (Path): Path to the output DOCX file.
        df (pd.DataFrame): DataFrame containing references with 'OpenAlexID' column.
        offset (int): Starting index for numbering references.
    Returns:
        int: The next offset after listing the references.
    """ 
    pyalex_ids = df['OpenAlexID'].dropna().unique().tolist()
    refstrings = tqdm.tqdm(create_refstrings_list(pyalex_ids))
    document = Document()
    for refstring in sorted(refstrings):
        p = document.add_paragraph()
        p.add_run(f"{offset}\t").bold = True
        p.add_run(refstring)
        offset += 1
    document.save(output_file)
    return offset


def list_dataframe_rows(
        output_file: Path,
        df: pd.DataFrame,
        format_function: callable,
        offset: int = 1,
        document: Document = None
    ) -> int:
    """List rows from a DataFrame to a DOCX file.
    Args:
        output_file (Path): Path to the output DOCX file.
        df (pd.DataFrame): DataFrame containing rows with column values to build reference strings.
        format_function (callable, optional): Function to format a row into a reference string.
        offset (int): Starting index for numbering rows.
        document (Document, optional): Existing DOCX document to append to.
    Returns:
        int: The next offset after listing the rows.
    """
    doc = Document() if document is None else document
    for _, row in df.iterrows():
        p = doc.add_paragraph()
        p.add_run(f"{offset}\t").bold = True
        p.add_run(format_function(row))
        offset += 1
    if document is None:
        doc.save(output_file)
    return offset


def convert_references_to_docx(
        csv_file: Path, 
        output_file: Path, 
    ) -> int:
    """Convert a CSV file of references to a DOCX file.
    Args:
        csv_file (Path): Path to the input CSV file.
        output_file (Path): Path to the output DOCX file.
    Returns:
        int: The number of references processed.
    """
    df = pd.read_csv(csv_file)
    document = Document()
    for id, refstr in df.itertuples(index=False):
        p = document.add_paragraph()
        p.add_run(f"{id}\t").bold = True
        p.add_run(refstr)
    document.save(output_file)
    return offset


# Main script execution
args = parse_arguments()
input_dir = Path(args.input_dir)
if not input_dir.is_dir():
    print(f"Error: The specified input directory '{input_dir}' does not exist or is not a directory.")
    sys.exit(1)
output_dir = Path(args.output_dir)
output_dir.mkdir(parents=True, exist_ok=True)

# Process FTE data
fte_file = input_dir / "example_fte_data.csv"
if os.path.exists(fte_file):
    df_fte = pd.read_csv(fte_file)
    color_variable = 'Activity'
    value_variable = 'Expenditure (FTE)'
    df_long = df_fte.melt(id_vars=['Year'], var_name=color_variable, value_name=value_variable)
    df_long['Activity'] = df_long['Activity'].apply(lambda x: '@'.join(wrap(x, 16)))
    chart = create_yearly_stacked_bar_chart(
        df=df_long,
        y_variable=value_variable,
        color_variable=color_variable,
        title="FTE Allocation per Activity",
        dimensions=[800, 500]
    )
    chart.save(output_dir / f"fte_allocation_chart.{args.format}")


# Make map of spatial collaborations diversity
cities_file = input_dir / "dutch_cities.csv"
if os.path.exists(cities_file):
    cities_df = pd.read_csv(cities_file)
    geo_chart = plot_netherlands_with_institutions(cities_df)
    geo_chart.save(f"netherlands_institutions.{args.format}")

# Make bar charts for funding data
funding_file = input_dir / "example_income_streams.csv"
if os.path.exists(funding_file):
    funding_df = pd.read_csv(funding_file)
    color_variable = 'Income Stream'
    value_variable = 'Income (M€)'
    df_long = funding_df.melt(id_vars=['Year'], var_name=color_variable, value_name=value_variable)
    chart = create_yearly_stacked_bar_chart(
        df=df_long,
        y_variable=value_variable,
        color_variable=color_variable,
        title="Income sources per Year",
        dimensions=[800, 500]
    )
    chart.save(output_dir / f"funding_data_chart.{args.format}")

# Make bar charts for headcount data
headcount_file = input_dir / "example_headcount_data.csv"
if os.path.exists(headcount_file):
    headcount_df = pd.read_csv(headcount_file)
    color_variable = 'Department'
    value_variable = 'Number of Employees'
    df_long = headcount_df.melt(id_vars=['Year'], var_name=color_variable, value_name=value_variable)
    chart = create_yearly_stacked_bar_chart(
        df=df_long,
        y_variable=value_variable,
        color_variable=color_variable,
        title="Number of Employees",
        dimensions=[800, 500]
    )
    chart.save(output_dir / f"headcount_data_chart.{args.format}")


# Make bar-line charts for calls data
calls_file = input_dir / "callData.csv"
if os.path.exists(calls_file):
    calls_df = pd.read_csv(calls_file, delimiter='|', encoding='utf-8')
    value_variable = 'Call budget (MEUR)'
    df_yearly_sums = calls_df.groupby(['Year'], as_index=False).sum()
    # Define constants
    acceptance_rate_variable = 'Acceptance Rate (%)'
    df_yearly_sums[acceptance_rate_variable] = 100 * df_yearly_sums['Acceptances'] / df_yearly_sums['Submissions']
    # Aggregate over 'Year' column and drop the 'Call' column
    calls_df['Call'] = calls_df['Call'].apply(lambda x: 'ASDI/OEC/ETEC' if x in ['ASDI', 'OEC', 'ETEC'] else 'Other')
    df_aggregated = calls_df.groupby(['Year', 'Call'], as_index=False).sum()
    df_aggregated[acceptance_rate_variable] = df_aggregated['Year'].map(df_yearly_sums.set_index('Year')[acceptance_rate_variable])
    bar_line_chart = create_yearly_stacked_bar_line_chart(
        df=df_aggregated,
        title="Submissions and Acceptance Rates",
        y_variable_left="Submissions",
        y_variable_right=acceptance_rate_variable,
        color_variable="Call",
        dimensions=[800, 500]
    )
    bar_line_chart.save(output_dir / f"calls_bar_line_chart.{args.format}")

    # Add 'Requested budget' and 'Provided budget' columns
    calls_df['Requested'] = (calls_df['In-kind funding per project (kEUR)'] + calls_df['Cash funding per project (kEUR)']) * calls_df['Submissions'] / 1000
    calls_df['Provided'] = (calls_df['In-kind funding per project (kEUR)'] + calls_df['Cash funding per project (kEUR)']) * calls_df['Acceptances'] /1000

    # Aggregate over 'Year' column and drop the 'Call' column
    df_aggregated = calls_df.groupby('Year', as_index=False).sum()
    df_aggregated.set_index('Year', inplace=True)
    df_aggregated = df_aggregated[['Requested', 'Provided']]

    offset_variable = 'Budget'
    df_long = df_aggregated.reset_index().melt(id_vars=['Year'], var_name=offset_variable, value_name=value_variable)
    chart = create_yearly_multi_bar_chart(
        df=df_long,
        title="Requested vs. Provided Budget",
        y_variable=value_variable,
        offset_variable=offset_variable,
        dimensions=[800, 500]
    )
    chart.save(output_dir / f"calls_multi_bar_chart.{args.format}")

    # Make ordered stacked bar chart of applicants
    institute_columns_applied = [col for col in calls_df.columns if col.endswith('Submissions') and col not in ['Submissions', 'Male Submissions', 'Female Submissions', 'NSE Submissions', 'EnvSus Submissions', 'LS Submissions', 'SSH Submissions']]
    institute_columns_granted = [col for col in calls_df.columns if col.endswith('Granted')]
    for col_a, col_g in zip(institute_columns_applied, institute_columns_granted):
        if str(col_a).replace("Submissions", "Granted") != col_g:
            print(f"Error: Column mismatch between {col_a} and {col_g}")
            continue
        calls_df[col_a] -= calls_df[col_g]
        calls_df[col_a] *= (calls_df['In-kind funding per project (kEUR)'] + calls_df['Cash funding per project (kEUR)']) / 1000.
        calls_df[col_g] *= (calls_df['In-kind funding per project (kEUR)'] + calls_df['Cash funding per project (kEUR)']) / 1000.

    df_aggregated = calls_df[institute_columns_applied + institute_columns_granted].sum().reset_index()
    df_aggregated.columns = ['Institute', BUDGET_COLUMN]
    df_aggregated["Category"] = ""
    df_aggregated["Category"] = df_aggregated["Institute"].apply(lambda x: "Applied" if x.endswith("Submissions") else "Granted")
    df_aggregated["Institute"] = df_aggregated["Institute"].str.replace(" Submissions", "")
    df_aggregated["Institute"] = df_aggregated["Institute"].str.replace(" Granted", "")
    df_aggregated["Institute"] = df_aggregated["Institute"].str.replace(" Submissions", "")
    chart = create_sorted_bar_chart(
        df=df_aggregated.reset_index(),
        category_variable='Institute',
        value_variable=BUDGET_COLUMN,
        title='Share of requested budget per institute',
        dimensions=[600, 600]
    )
    chart.save(output_dir / f"calls_institutes_pie_chart.{args.format}")

projects_file = input_dir / "projectOverview.csv"

if os.path.exists(projects_file):
    ext_proj_df = pd.read_csv(projects_file, index_col=0, delimiter='|', encoding='utf-8')
    # Reset index to make year a column if it's currently the index
    if ext_proj_df.index.name is None:
        ext_proj_df.index.name = 'slug'
    
    # Aggregate data by 'call_year' and 'TYPE'
    df_aggregated = ext_proj_df.groupby(['call_year', 'TYPE'], as_index=False)['INCOME'].sum()
    df_aggregated['INCOME'] = df_aggregated['INCOME'] / 1000  # Convert to kEUR

    # Rename columns for clarity
    df_aggregated.rename(columns={'call_year': 'Year', 'TYPE': 'Type', 'INCOME': 'Income (kEUR)'}, inplace=True)

    external_acquisition_chart = create_yearly_stacked_bar_chart(
        df=df_aggregated,
        y_variable="Income (kEUR)",
        color_variable="Type",
        title="External Acquisition Income",
        dimensions=[800, 500]
    )
    external_acquisition_chart.save(output_dir / f"external_acquisition_per_year.{args.format}")

def load_and_filter_publications(file_path: Path, filter_external: bool = True) -> pd.DataFrame:
    """Load publications CSV and filter out EXTERNAL projects with no author position."""
    df = pd.DataFrame()
    if os.path.exists(file_path):
        df = pd.read_csv(file_path, delimiter='|', encoding='utf-8')
        if filter_external:
            df = df[~((df['project type'] == 'EXTERNAL') & (df['author position'] == 'none'))]
    return df

journal_df = load_and_filter_publications(input_dir / "journalArticle.csv")
conference_df = load_and_filter_publications(input_dir / "conferencePaper.csv")
book_df = load_and_filter_publications(input_dir / "book.csv")
publications_df = pd.concat([journal_df, conference_df, book_df], ignore_index=True)

if not publications_df.empty:
    # Create new column 'authorship' based on 'author position'
    publications_df['authorship'] = publications_df['author position'].apply(lambda x: 'NLeSC funded' if x == 'none' else 'NLeSC authored')
    # Aggregate data by 'year' and 'authorship'
    df_aggregated = publications_df.groupby(['year', 'authorship'], as_index=False).size()
    df_aggregated.rename(columns={'year': 'Year', 'size': 'Number of Publications'}, inplace=True)
    publications_chart = create_yearly_stacked_bar_chart(
        df=df_aggregated,
        y_variable="Number of Publications",
        color_variable="authorship",
        title="Peer-reviewed publications per year",
        dimensions=[800, 500]
    )
    publications_chart.save(output_dir / f"publications_per_year.{args.format}")
    total_publications = df_aggregated['Number of Publications'].sum()
    total_authored_publications = df_aggregated['Number of Publications'][df_aggregated['authorship'] == 'NLeSC authored'].sum()
    total_first_author_publications = publications_df[publications_df['author position'] == 'first'].shape[0]
    total_citations = publications_df['citations'].sum()
    total_authored_citations = publications_df[publications_df['authorship'] == 'NLeSC authored']['citations'].sum()
    total_open_access = publications_df['open access'].sum()
    print(f"Total publications: {total_publications}")
    print(f"Total journal publications: {journal_df.shape[0]}")
    print(f"Total conference publications: {conference_df.shape[0]}")
    print(f"Total book publications: {book_df.shape[0]}")
    print(f"Total NLeSC authored publications: {total_authored_publications}")
    print(f"Total first-author publications: {total_first_author_publications}")
    print(f"Total citations: {total_citations}")
    print(f"Total citations for NLeSC authored publications: {total_authored_citations}")
    print(f"Open access publications: {total_open_access} ({(total_open_access/total_publications)*100:.2f}%)")

    offset = 1
    # Generate reference lists
    papers_df = pd.concat([journal_df, conference_df], ignore_index=True)
    papers_df['authorship'] = papers_df['author position'].apply(lambda x: 'NLeSC funded' if x == 'none' else 'NLeSC authored')

    authored_papers_file = output_dir / "authored_papers.docx"
    if not authored_papers_file.exists():
        offset = list_references(authored_papers_file, papers_df[papers_df['authorship'] == 'NLeSC authored'], offset)
    else:
        offset += papers_df[papers_df['authorship'] == 'NLeSC authored'].shape[0]
    funded_papers_file = output_dir / "funded_papers.docx"
    if not funded_papers_file.exists():
        offset = list_references(funded_papers_file, papers_df[papers_df['authorship'] == 'NLeSC funded'], offset)
    else:        
        offset += papers_df[papers_df['authorship'] == 'NLeSC funded'].shape[0]
    books_file = output_dir / "books.docx"
    if not books_file.exists():
        offset = list_references(books_file, book_df, offset)
    else:
        offset += book_df.shape[0]

# Generate reference lists for preprints and reports
preprint_df = load_and_filter_publications(input_dir / "preprint.csv")
report_df = load_and_filter_publications(input_dir / "report.csv")
publications_df = pd.concat([preprint_df, report_df], ignore_index=True)
if not publications_df.empty:
    # Create new column 'authorship' based on 'author position'
    publications_df['authorship'] = publications_df['author position'].apply(lambda x: 'NLeSC funded' if x == 'none' else 'NLeSC authored')
    total_authored_publications = publications_df[publications_df['authorship'] == 'NLeSC authored'].shape[0]
    total_first_author_publications = publications_df[publications_df['author position'] == 'first'].shape[0]
    print(f"Total preprints and reports: {publications_df.shape[0]}")
    print(f"Total NLeSC authored preprints and reports: {total_authored_publications}")
    print(f"Total first-author preprints and reports: {total_first_author_publications}")

    # Generate reference lists
    white_papers_file = output_dir / "white_papers.docx"
    if not white_papers_file.exists():
        offset = list_references(white_papers_file, publications_df, offset)
    else:        
        offset += publications_df.shape[0]

# Generate reference lists for press releases
press_releases_df = load_and_filter_publications(input_dir / "press.csv")
if not press_releases_df.empty:
    print(f"Total press releases: {press_releases_df.shape[0]}")
    press_releases_file = output_dir / "press_releases.docx"
    if not press_releases_file.exists():
        offset = list_dataframe_rows(press_releases_file, press_releases_df, 
                                     format_function=lambda row: f"{row['title']} ({row['year']}), {row.get('url', '')}",
                                     offset=offset)
    else:
        offset += press_releases_df.shape[0]

# Generate reference lists data publications
data_publications_df = load_and_filter_publications(input_dir / "data.csv")
if not data_publications_df.empty:
    print(f"Total data publications: {data_publications_df.shape[0]}")
    data_publications_file = output_dir / "data_publications.docx"
    if not data_publications_file.exists():
        offset = list_references(data_publications_file, data_publications_df, offset)
    else:
        offset += data_publications_df.shape[0]

# Generate reference lists for blog posts
blog_posts_df = load_and_filter_publications(input_dir / "blog.csv", filter_external=False)
if not blog_posts_df.empty:
    print(f"Total blog posts: {blog_posts_df.shape[0]}")
    blog_posts_file = output_dir / "blog_posts.docx"
    if not blog_posts_file.exists():
        offset = list_dataframe_rows(blog_posts_file, blog_posts_df, 
                                    format_function=lambda row: f"{row['author']}. {row['title']} ({pd.to_datetime(row['date']).year}), {row.get('url', '')}",
                                    offset=offset)
    else:
        offset += blog_posts_df.shape[0]

# Generate reference lists for theses
theses_df = load_and_filter_publications(input_dir / "thesis.csv", filter_external=False)
if not theses_df.empty:
    phd_theses_df = theses_df[theses_df["type"] == "PhD thesis"]
    print(f"Total PhD theses: {phd_theses_df.shape[0]}")
    master_theses_df = theses_df[theses_df["type"] == "Master thesis"]
    print(f"Total master theses: {master_theses_df.shape[0]}")
    bachelor_theses_df = theses_df[theses_df["type"] == "Bachelor thesis"]
    print(f"Total bachelor theses: {bachelor_theses_df.shape[0]}")
    theses_file = output_dir / "theses.docx"
    if not theses_file.exists():
        reference_builder = lambda row: f"{row['author']}. {row['title']}, {row['source']}, {row['type']} ({row['year']}). {row.get('url', '')}"
        document = Document()
        offset = list_dataframe_rows(theses_file, phd_theses_df, format_function=reference_builder, offset=offset, document=document)
        offset = list_dataframe_rows(theses_file, master_theses_df, format_function=reference_builder,offset=offset, document=document)
        offset = list_dataframe_rows(theses_file, bachelor_theses_df, format_function=reference_builder, offset=offset, document=document)
        document.save(theses_file)
    else:
        offset += theses_df.shape[0]

# Generate reference lists for workshops
workshops_df = load_and_filter_publications(input_dir / "workshop.csv", filter_external=False)
if not workshops_df.empty:
    print(f"Total workshops: {workshops_df.shape[0]}")
    workshops_file = output_dir / "workshops.docx"
    if not workshops_file.exists():
        offset = list_dataframe_rows(workshops_file, workshops_df, 
                                    format_function=lambda row: f"{row['title']} ({row['year']}), {row.get('url', '')}",
                                    offset=offset)
    else:
        offset += workshops_df.shape[0]


LA_surveys_file = input_dir / "LASurveyScores.csv"
if os.path.exists(LA_surveys_file):
    la_survey_df = pd.read_csv(LA_surveys_file, delimiter='|', encoding='utf-8')
    for subset in la_survey_df.groupby('Section'):
        section_name, section_df = subset
        section_df_long = section_df.melt(id_vars=['Question'], 
                                            value_vars=['average', '2018 average'],
                                            var_name="period", value_name="score")
        replacement_values = {"average": "2019-2025", "2018 average": "2013-2018"}
        section_df_long["period"] = section_df_long["period"].replace(replacement_values)
        section_df_long['Question'] = section_df_long['Question'].apply(lambda x: '@'.join(wrap(x, 40)))
        num_questions = section_df_long['Question'].nunique()
        chart = create_survey_chart(
            df=section_df_long,
            x_variable='score',
            x_variable_2='period',
            dimensions=[800, 160 * num_questions],
        )
        chart.save(output_dir / f"LA_survey_{section_name}.{args.format}")

training_data_file = input_dir / "trainingData.csv"
if os.path.exists(training_data_file):
    training_df = pd.read_csv(training_data_file, delimiter='|', encoding='utf-8')
    training_df_long = training_df.drop(columns=["Topic"]).melt(id_vars="Indicator", var_name="Year").pivot_table(columns="Indicator", values='value', index='Year').reset_index()
    chart = create_yearly_multi_line_chart( training_df_long, 
                                            y_variable_left="Subscriptions",
                                            y_variable_right="Attendance rate (%)",
                                            title="eScience Trainings Demand",
                                            charge_year="2024",
                                            y_scale_left=[0, 700])
    chart.save(output_dir / f"training_attendance.{args.format}")

    training_df_regional = training_df[training_df["Topic"] == "Attendees region"].drop(columns=["Topic"]).melt(id_vars="Indicator", var_name="Year").pivot_table(columns="Indicator", values='value', index='Year').reset_index()
    training_df_year_average = training_df_regional.mean(numeric_only=True, axis=0).reset_index()
    training_df_year_average.columns = ["Region", "Attendees"]
    training_df_year_average['Region'] = training_df_year_average['Region'].str.replace('Attendees ', '', regex=False)
    chart = create_pie_chart(   training_df_year_average, 
                                title="Average attendance per region",
                                category_variable="Region",
                                value_variable="Attendees")
    chart.save(output_dir / f"training_attendance_region.{args.format}")

    training_df_career_stage = training_df[training_df["Topic"] == "Attendees domain"].drop(columns=["Topic"]).melt(id_vars="Indicator", var_name="Year").pivot_table(columns="Indicator", values='value', index='Year').reset_index()
    training_df_career_stage_year_average = training_df_career_stage.mean(numeric_only=True, axis=0).reset_index()
    training_df_career_stage_year_average.columns = ["Domain", "Attendees"]
    training_df_career_stage_year_average['Domain'] = training_df_career_stage_year_average['Domain'].str.replace('Attendees ', '', regex=False)
    domain_mapping = {'LSH': 'Life Sciences', 
                      'SSH': 'Social Sciences@and Humanities',
                      'NES': 'Natural Sciences@and Engineering', 
                      'ENV': 'Environment@and Sustainability'}
    training_df_career_stage_year_average['Domain'] = training_df_career_stage_year_average['Domain'].map(domain_mapping)

    chart = create_pie_chart(   training_df_career_stage_year_average, 
                                title="Average attendance per domain",
                                category_variable="Domain",
                                value_variable="Attendees")
    chart.save(output_dir / f"training_attendance_domain.{args.format}")

    training_df_survey = training_df[training_df["Topic"] == "Survey"].drop(columns=["Topic"]).melt(id_vars="Indicator", var_name="Year").reset_index()
    training_df_survey['Indicator'] = training_df_survey['Indicator'].apply(lambda x: '@'.join(wrap(x, 40)))
    chart = create_table_heatmap(   training_df_survey, 
                                    title="Training survey results",
                                    x_variable="Year",
                                    y_variable="Indicator",
                                    value_variable="value",
                                    x_range=[2022, 2026])
    chart.save(output_dir / f"training_survey.{args.format}")





